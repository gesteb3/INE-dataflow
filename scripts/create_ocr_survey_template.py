"""Genera la plantilla imprimible que usa el MVP de OCR."""

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "data" / "templates" / "encuesta_hogares_ocr_ine.docx"
BLUE = RGBColor(31, 77, 120)
LIGHT_BLUE = "E8EEF5"
GRAY = RGBColor(90, 90, 90)


def set_cell_shading(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    properties.append(shading)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    properties = table._tbl.tblPr
    width = OxmlElement("w:tblW")
    width.set(qn("w:w"), str(sum(widths)))
    width.set(qn("w:type"), "dxa")
    properties.append(width)
    indent = OxmlElement("w:tblInd")
    indent.set(qn("w:w"), "120")
    indent.set(qn("w:type"), "dxa")
    properties.append(indent)
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for value in widths:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(value))
        grid.append(column)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = Inches(widths[index] / 1440)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell_width = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if cell_width is None:
                cell_width = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(cell_width)
            cell_width.set(qn("w:w"), str(widths[index]))
            cell_width.set(qn("w:type"), "dxa")


def style_run(run, size=11, color=None, bold=False, italic=False):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color


def add_field_row(table, left_label, left_value, right_label, right_value):
    row = table.add_row()
    values = [left_label, left_value, right_label, right_value]
    for cell, value in zip(row.cells, values):
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(value)
        style_run(run, size=10, color=BLUE if cell is row.cells[0] or cell is row.cells[2] else GRAY, bold=cell is row.cells[0] or cell is row.cells[2])
        if cell is row.cells[0] or cell is row.cells[2]:
            set_cell_shading(cell, LIGHT_BLUE)


def add_section_heading(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(14)
    paragraph.paragraph_format.space_after = Pt(7)
    run = paragraph.add_run(text)
    style_run(run, size=13, color=BLUE, bold=True)


def main():
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    style_run(header.add_run("INE DataFlow · Plantilla OCR v1.0"), size=8, color=GRAY)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_run(footer.add_run("Documento académico · Usar letra de molde y tinta oscura"), size=8, color=GRAY)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(3)
    style_run(title.add_run("ENCUESTA DE HOGARES"), size=22, color=BLUE, bold=True)
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(12)
    style_run(subtitle.add_run("Formulario estandarizado para captura mediante OCR"), size=12, color=GRAY)

    note = document.add_paragraph()
    note.paragraph_format.space_after = Pt(8)
    style_run(note.add_run("Instrucciones: "), size=10, color=BLUE, bold=True)
    style_run(note.add_run("complete los campos con letra de molde, una respuesta por casilla y sin tachones. Para la demostración se recomienda escribir en MAYÚSCULAS."), size=10, color=GRAY)

    add_section_heading(document, "1. Identificación de la encuesta")
    table = document.add_table(rows=0, cols=4)
    set_table_geometry(table, [1850, 2830, 1850, 2830])
    add_field_row(table, "Código de encuesta", "________________", "Fecha (AAAA-MM-DD)", "____-__-__")
    add_field_row(table, "Departamento", "______________", "Municipio", "______________")

    add_section_heading(document, "2. Características del hogar")
    table = document.add_table(rows=0, cols=4)
    set_table_geometry(table, [1850, 2830, 1850, 2830])
    add_field_row(table, "Área", "☐ URBANA   ☐ RURAL", "Personas en hogar", "________")
    add_field_row(table, "Edad (años)", "________", "Sexo", "☐ F   ☐ M   ☐ O")
    add_field_row(table, "Ingreso mensual GTQ", "________________", "Visita", "☐ 1   ☐ 2   ☐ 3")

    add_section_heading(document, "3. Observaciones del encuestador")
    for _ in range(3):
        paragraph = document.add_paragraph("________________________________________________________________________________")
        paragraph.paragraph_format.space_after = Pt(5)
        style_run(paragraph.runs[0], size=10, color=GRAY)

    note = document.add_paragraph()
    note.paragraph_format.space_before = Pt(12)
    note.paragraph_format.space_after = Pt(0)
    style_run(note.add_run("Campos que reconoce el OCR: "), size=9, color=BLUE, bold=True)
    style_run(note.add_run("survey_code, interview_date, department_code, municipality_code, urban_rural, respondent_age, respondent_sex, household_size y monthly_income_gtq."), size=9, color=GRAY)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
